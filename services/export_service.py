"""
services/export_service.py
==========================================
Ekspor tabel asuhan ke Markdown, Word (.docx), dan Excel (.xlsx).

BENTUK TABEL
------------
Ketiganya memakai tata letak yang sama, mengikuti format lembar asuhan
keperawatan yang lazim dipakai:

    | No | Diagnosis Keperawatan (SDKI) | Luaran (SLKI) | Intervensi (SIKI) |

Satu baris = satu diagnosis. Intervensi ditulis di dalam satu sel,
dikelompokkan per kategori SIKI dengan penomoran.

Kenapa bukan satu baris per tindakan: bentuk itu memang lebih rapi untuk
diolah sebagai data, tetapi tidak menyerupai lembar askep yang dipakai
sehari-hari dan tidak bisa langsung ditempel. Karena tujuan ekspor ini
untuk disalin ke dokumen asuhan, bentuk lembarlah yang dimenangkan.

Ekspor CSV versi lama dihapus. CSV menyatukan seluruh intervensi menjadi
satu untaian panjang dipisah titik koma — pada D.0008 hasilnya 344
karakter menggumpal dalam satu sel, praktis tidak terbaca. Berkas .xlsx
menggantikannya dengan sel multi-baris, pembungkusan teks, tinggi baris
yang menyesuaikan, dan lebar kolom yang sudah diatur.
"""

from __future__ import annotations

import io
from typing import Any

from models.asesmen import Asesmen

_KATEGORI_SIKI = ("observasi", "terapeutik", "edukasi", "kolaborasi")
_LABEL_SIKI = {
    "observasi": "Observasi",
    "terapeutik": "Terapeutik",
    "edukasi": "Edukasi",
    "kolaborasi": "Kolaborasi",
}

JUDUL_KOLOM = [
    "No",
    "Diagnosis Keperawatan (SDKI)",
    "Luaran (SLKI)",
    "Intervensi Keperawatan (SIKI)",
]


# =====================================================
# PENYUSUN ISI SEL
# =====================================================

def _tindakan_terpilih(item: dict[str, Any], kategori: str) -> list[str]:
    """
    Tindakan yang ditampilkan untuk satu kategori.

    Bila perawat tidak mencentang apa pun, SELURUH intervensi ditampilkan
    sebagai acuan — lebih berguna daripada kolom kosong bagi yang
    melewati langkah pemilihan.
    """
    semua = (item.get("intervensi") or {}).get(kategori) or []
    dipilih = set(item.get("intervensi_dipilih") or [])
    if not dipilih:
        return list(semua)
    return [t for t in semua if t in dipilih]


def _sel_diagnosis(item: dict[str, Any]) -> str:
    baris = [item.get("kode", ""), item.get("nama", "")]
    keterangan = " · ".join(x for x in (item.get("jenis"), item.get("kategori")) if x)
    if keterangan:
        baris.append(f"({keterangan})")
    return "\n".join(b for b in baris if b)


def _sel_luaran(item: dict[str, Any]) -> str:
    """
    Kolom luaran sebagai SATU KESATUAN kalimat askep.

    Mengikuti rumusan yang lazim dipakai di lembar asuhan keperawatan:

        Setelah dilakukan intervensi keperawatan selama 8 jam,
        maka Curah Jantung Meningkat, dengan kriteria hasil:
        1. Kekuatan nadi perifer 4-5 (awal: ....)
        2. Frekuensi nadi 60-100 x/menit (awal: ....)

    Versi sebelumnya memisahkan luaran, waktu evaluasi, dan daftar
    indikator menjadi beberapa blok. Bentuk itu lebih rapi untuk dibaca
    sebagai data, tetapi bukan bentuk yang bisa langsung disalin ke
    lembar askep — sedangkan justru itulah tujuan ekspor ini.

    Baseline tetap disediakan sebagai isian titik-titik dalam kurung,
    bukan kolom tersendiri: nilai awal harus berasal dari penilaian
    perawat, dan menaruhnya inline membuat kalimatnya tetap utuh.
    """
    luaran = item.get("luaran") or {}
    kode = luaran.get("kode", "")
    nama = luaran.get("nama", "")

    baris: list[str] = []
    if kode or nama:
        baris.append(" — ".join(x for x in (kode, nama) if x))

    indikator = item.get("indikator") or []
    if not indikator:
        return "\n".join(baris)

    waktu = item.get("evaluasi_default") or "24 jam"
    baris.append("")
    baris.append(
        f"Setelah dilakukan intervensi keperawatan selama {waktu}, "
        f"maka {nama or 'luaran tercapai'}, dengan kriteria hasil:"
    )

    for nomor, i in enumerate(indikator, start=1):
        satuan = f" {i['satuan']}" if i.get("satuan") and i["satuan"] != "—" else ""
        target = i.get("target", "")
        baris.append(f"{nomor}. {i['nama']} {target}{satuan} (awal: ....)")

    return "\n".join(baris)


def _blok_intervensi(item: dict[str, Any]) -> list[tuple[str, list[str]]]:
    """[(label kategori, [tindakan, ...]), ...] — kategori kosong dibuang."""
    blok = []
    for kategori in _KATEGORI_SIKI:
        tindakan = _tindakan_terpilih(item, kategori)
        if tindakan:
            blok.append((_LABEL_SIKI[kategori], tindakan))
    return blok


def _sel_intervensi(item: dict[str, Any]) -> str:
    bagian = []
    for label, tindakan in _blok_intervensi(item):
        baris = [label]
        baris += [f"{i}. {t}" for i, t in enumerate(tindakan, start=1)]
        bagian.append("\n".join(baris))
    return "\n\n".join(bagian) or "-"


def _judul(asesmen: Asesmen) -> str:
    bagian = ["Rencana Asuhan Keperawatan"]
    if asesmen.nomor:
        bagian.append(asesmen.nomor)
    return " — ".join(bagian)


def nama_berkas(asesmen: Asesmen, ekstensi: str) -> str:
    dasar = asesmen.nomor or "asuhan"
    return f"{dasar}.{ekstensi.lstrip('.')}"


# =====================================================
# MARKDOWN
# =====================================================

def ke_markdown(asesmen: Asesmen, tabel: list[dict[str, Any]]) -> str:
    baris: list[str] = [f"# {_judul(asesmen)}", ""]
    if asesmen.label:
        baris.append(f"**Penanda:** {asesmen.label}  ")
    baris.append(f"**Dibuat:** {asesmen.dibuat_pada or '-'}")
    baris += ["", "## Data Asesmen", ""]
    baris.append(f"**S (Subjektif):** {asesmen.data_subjektif or '-'}")
    baris += ["", f"**O (Objektif):** {asesmen.data_objektif or '-'}", ""]

    if not tabel:
        baris.append("_Belum ada diagnosis yang dipilih._")
        return "\n".join(baris)

    baris += ["## Diagnosis, Luaran, dan Intervensi", ""]

    for item in tabel:
        luaran = item.get("luaran") or {}
        baris.append(f"### {item['prioritas']}. {item['kode']} — {item['nama']}")
        baris.append("")
        baris.append(f"- **Jenis:** {item.get('jenis') or '-'}")
        baris.append(f"- **Kategori:** {item.get('kategori') or '-'}")
        baris.append(
            f"- **Luaran (SLKI):** {luaran.get('kode', '-')} — {luaran.get('nama', '-')}"
        )
        # Sisipkan kriteria hasil dari penyusun yang sama dengan Word dan
        # Excel, supaya ketiga format tidak berbeda isi.
        indikator = item.get("indikator") or []
        if indikator:
            waktu = item.get("evaluasi_default") or "24 jam"
            baris.append("")
            baris.append(
                f"Setelah dilakukan intervensi keperawatan selama **{waktu}**, maka "
                f"**{luaran.get('nama', '')}**, dengan kriteria hasil:"
            )
            baris.append("")
            for nomor, i in enumerate(indikator, start=1):
                satuan = f" {i['satuan']}" if i.get("satuan") and i["satuan"] != "—" else ""
                baris.append(f"{nomor}. {i['nama']} **{i.get('target', '')}**{satuan} (awal: ....)")
        if item.get("catatan"):
            baris.append(f"- **Catatan:** {item['catatan']}")
        baris += ["", "**Intervensi (SIKI)**", ""]

        for label, tindakan in _blok_intervensi(item):
            baris.append(f"*{label}*")
            baris += [f"{i}. {t}" for i, t in enumerate(tindakan, start=1)]
            baris.append("")
        baris.append("")

    if asesmen.catatan:
        baris += ["## Catatan", "", asesmen.catatan, ""]

    return "\n".join(baris)


# =====================================================
# WORD (.docx)
# =====================================================

def ke_docx(asesmen: Asesmen, tabel: list[dict[str, Any]]) -> bytes:
    """
    Dokumen Word berisi tabel askep siap salin-tempel.

    Mengembalikan bytes agar bisa langsung dipakai `st.download_button`.
    """
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "Ekspor Word butuh python-docx. Install dengan: pip install python-docx"
        ) from exc

    dok = Document()

    # Halaman A4 lanskap. Ukuran ditetapkan eksplisit karena bawaan
    # python-docx adalah Letter (216 x 279 mm) — di Indonesia yang dipakai
    # A4, dan selisihnya membuat tabel tidak pas saat dicetak.
    # Lanskap dipilih karena tabel empat kolom dengan daftar intervensi
    # terlalu sempit pada halaman potret.
    bagian = dok.sections[0]
    bagian.orientation = WD_ORIENT.LANDSCAPE
    bagian.page_width = Pt(842)   # A4 sisi panjang
    bagian.page_height = Pt(595)  # A4 sisi pendek
    for sisi in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(bagian, sisi, Pt(36))

    lebar_tersedia = bagian.page_width.pt - bagian.left_margin.pt - bagian.right_margin.pt

    judul = dok.add_heading(_judul(asesmen), level=1)
    judul.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = []
    if asesmen.label:
        meta.append(f"Penanda: {asesmen.label}")
    if asesmen.dibuat_pada:
        meta.append(f"Dibuat: {asesmen.dibuat_pada}")
    if meta:
        p = dok.add_paragraph("   |   ".join(meta))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    dok.add_heading("Data Asesmen", level=2)
    for label, isi in (
        ("S (Subjektif)", asesmen.data_subjektif),
        ("O (Objektif)", asesmen.data_objektif),
    ):
        p = dok.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(isi or "-")

    dok.add_heading("Rencana Asuhan", level=2)

    if not tabel:
        dok.add_paragraph("Belum ada diagnosis yang dipilih.")
        keluaran = io.BytesIO()
        dok.save(keluaran)
        return keluaran.getvalue()

    t = dok.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Tanpa tiga baris berikut, Word MENGABAIKAN lebar kolom yang diset dan
    # melebarkan tabel mengikuti isi terpanjang — kolom intervensi meluber
    # keluar halaman dan teksnya terpotong saat dicetak.
    #
    # `autofit = False` saja tidak cukup; Word baru benar-benar mematuhi
    # lebar kolom bila algoritma tata letaknya dikunci ke "fixed" lewat
    # elemen tblLayout, dan lebar tabel keseluruhan dinyatakan di tblW.
    t.autofit = False
    tbl_pr = t._tbl.tblPr

    tata_letak = OxmlElement("w:tblLayout")
    tata_letak.set(qn("w:type"), "fixed")
    tbl_pr.append(tata_letak)

    lebar_tabel = OxmlElement("w:tblW")
    lebar_tabel.set(qn("w:type"), "dxa")
    lebar_tabel.set(qn("w:w"), str(int(lebar_tersedia * 20)))  # dxa = 1/20 pt
    tbl_pr.append(lebar_tabel)

    kepala = t.rows[0].cells
    for idx, teks in enumerate(JUDUL_KOLOM):
        kepala[idx].text = ""
        p = kepala[idx].paragraphs[0]
        r = p.add_run(teks)
        r.bold = True
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for item in tabel:
        sel = t.add_row().cells

        sel[0].text = str(item.get("prioritas", ""))
        sel[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Kolom diagnosis: kode tebal, nama biasa, keterangan kecil & abu.
        sel[1].text = ""
        p = sel[1].paragraphs[0]
        p.add_run(item.get("kode", "")).bold = True
        p.add_run("\n" + (item.get("nama") or ""))
        keterangan = " · ".join(x for x in (item.get("jenis"), item.get("kategori")) if x)
        if keterangan:
            r = p.add_run("\n" + keterangan)
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Kolom luaran memakai _sel_luaran() yang sama dengan Excel dan
        # Markdown, supaya ketiga format menghasilkan isi yang identik.
        # Sebelumnya sel ini dibangun terpisah, sehingga indikator dan
        # waktu evaluasi tidak ikut tercetak di Word — perbedaan yang
        # tidak terlihat sampai berkasnya dibuka.
        luaran = item.get("luaran") or {}
        sel[2].text = ""
        p = sel[2].paragraphs[0]
        potongan = _sel_luaran(item).split("\n")
        for nomor, teks in enumerate(potongan):
            r = p.add_run(("\n" if nomor else "") + teks)
            # Baris pertama berisi kode dan nama luaran — ditebalkan agar
            # tetap menonjol di antara daftar kriteria hasil.
            if nomor == 0:
                r.bold = True
            else:
                r.font.size = Pt(9)

        # Kolom intervensi: satu paragraf per kategori + tindakan bernomor.
        sel[3].text = ""
        pertama = True
        for label, tindakan in _blok_intervensi(item):
            p = sel[3].paragraphs[0] if pertama else sel[3].add_paragraph()
            pertama = False
            r = p.add_run(label)
            r.bold = True
            r.font.size = Pt(9)
            for i, aksi in enumerate(tindakan, start=1):
                r = p.add_run(f"\n{i}. {aksi}")
                r.font.size = Pt(9)

        for c in sel:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)

    # Lebar kolom dihitung sebagai bagian dari lebar halaman, bukan angka
    # mati — supaya tetap pas bila ukuran kertas atau margin diubah.
    # Kolom intervensi mendapat porsi terbesar karena isinya paling panjang.
    porsi = (0.04, 0.22, 0.18, 0.56)
    lebar_kolom = [Pt(lebar_tersedia * x) for x in porsi]

    # Lebar harus diset pada SETIAP sel; Word mengabaikan lebar yang hanya
    # dinyatakan di level kolom.
    for baris in t.rows:
        for idx, c in enumerate(baris.cells):
            c.width = lebar_kolom[idx]

    if asesmen.catatan:
        dok.add_heading("Catatan", level=2)
        dok.add_paragraph(asesmen.catatan)

    keluaran = io.BytesIO()
    dok.save(keluaran)
    return keluaran.getvalue()


# =====================================================
# EXCEL (.xlsx)
# =====================================================

def ke_xlsx(asesmen: Asesmen, tabel: list[dict[str, Any]]) -> bytes:
    """Berkas Excel berisi tabel askep dengan sel multi-baris yang terbaca."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "Ekspor Excel butuh openpyxl. Install dengan: pip install openpyxl"
        ) from exc

    wb = Workbook()
    ws = wb.active
    ws.title = "Asuhan Keperawatan"

    garis = Side(style="thin", color="BFBFBF")
    kotak = Border(left=garis, right=garis, top=garis, bottom=garis)
    atas_kiri = Alignment(vertical="top", wrap_text=True)

    ws.append([_judul(asesmen)])
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=4)

    baris_meta = []
    if asesmen.label:
        baris_meta.append(f"Penanda: {asesmen.label}")
    if asesmen.dibuat_pada:
        baris_meta.append(f"Dibuat: {asesmen.dibuat_pada}")
    ws.append(["   |   ".join(baris_meta)])
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=4)
    ws["A2"].font = Font(size=9, color="666666")

    ws.append([])
    for label, isi in (
        ("S (Subjektif)", asesmen.data_subjektif),
        ("O (Objektif)", asesmen.data_objektif),
    ):
        ws.append([label, isi or "-"])
        r = ws.max_row
        ws[f"A{r}"].font = Font(bold=True)
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=4)
        ws[f"B{r}"].alignment = atas_kiri
    ws.append([])

    baris_judul = ws.max_row + 1
    ws.append(JUDUL_KOLOM)
    latar = PatternFill("solid", fgColor="0D9488")
    for idx in range(1, 5):
        sel = ws.cell(row=baris_judul, column=idx)
        sel.font = Font(bold=True, color="FFFFFF", size=11)
        sel.fill = latar
        sel.alignment = Alignment(vertical="center", horizontal="center", wrap_text=True)
        sel.border = kotak

    for item in tabel:
        ws.append([
            item.get("prioritas", ""),
            _sel_diagnosis(item),
            _sel_luaran(item),
            _sel_intervensi(item),
        ])
        r = ws.max_row
        for idx in range(1, 5):
            sel = ws.cell(row=r, column=idx)
            sel.alignment = atas_kiri
            sel.border = kotak
        ws.cell(row=r, column=1).alignment = Alignment(vertical="top", horizontal="center")

        # Tinggi baris diperkirakan dari isi terpanjang. Tanpa ini Excel
        # memakai tinggi satu baris dan hampir seluruh isi tersembunyi.
        jumlah_baris_teks = max(
            len(str(ws.cell(row=r, column=k).value or "").split("\n")) for k in (2, 3, 4)
        )
        ws.row_dimensions[r].height = min(max(jumlah_baris_teks * 13.5, 30), 409)

    for kolom, lebar in ((1, 5), (2, 32), (3, 28), (4, 78)):
        ws.column_dimensions[get_column_letter(kolom)].width = lebar

    ws.freeze_panes = ws.cell(row=baris_judul + 1, column=1)

    # Siapkan untuk dicetak: lanskap, muat selebar satu halaman.
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_title_rows = f"{baris_judul}:{baris_judul}"

    if asesmen.catatan:
        ws.append([])
        ws.append(["Catatan", asesmen.catatan])
        r = ws.max_row
        ws[f"A{r}"].font = Font(bold=True)
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=4)
        ws[f"B{r}"].alignment = atas_kiri

    keluaran = io.BytesIO()
    wb.save(keluaran)
    return keluaran.getvalue()


# =====================================================
# KETERSEDIAAN
# =====================================================

def docx_tersedia() -> bool:
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


def xlsx_tersedia() -> bool:
    try:
        import openpyxl  # noqa: F401

        return True
    except ImportError:
        return False
