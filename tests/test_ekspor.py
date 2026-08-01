"""
tests/test_ekspor.py
==========================================
Test ekspor tabel asuhan ke Markdown, Word, dan Excel.

LATAR BELAKANG
--------------
Ekspor CSV versi lama menyatukan seluruh intervensi menjadi satu untaian
dipisah titik koma. Pada D.0008 hasilnya 344 karakter menggumpal dalam
satu sel — secara teknis "berhasil", tetapi tidak terbaca dan tidak bisa
dipakai. CSV diganti .xlsx dan .docx yang berisi tabel askep sungguhan.

Test ini memeriksa hasilnya dengan MEMBUKA KEMBALI berkas yang dibuat,
bukan sekadar memastikan fungsinya tidak melempar error. Berkas yang
terbentuk tapi isinya kacau tetap kegagalan.

Jalankan:
    cd tests && python test_ekspor.py
"""

from __future__ import annotations

import io
import os
import sys
import tempfile
import types
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
os.environ["ASUHAN_DB_PATH"] = os.path.join(tempfile.mkdtemp(prefix="ekspor_test_"), "t.db")

_st = types.ModuleType("streamlit")
_st.secrets = {}
sys.modules.setdefault("streamlit", _st)

from models.asesmen import Asesmen, DiagnosisPilihan  # noqa: E402
from services import export_service as E  # noqa: E402
from services.diagnosis_service import DiagnosisService  # noqa: E402

PASS = 0
FAIL = 0


def check(label, cond, extra=""):
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  [PASS] {label}")
    else:
        FAIL += 1
        print(f"  [FAIL] {label}  {extra}")


def contoh() -> tuple[Asesmen, list]:
    asesmen = Asesmen(
        nomor="ASM-20260731-001",
        label="Bed 3",
        data_subjektif="Mengeluh sesak saat berbaring, mudah lelah",
        data_objektif="Edema tungkai, JVP meningkat, ronkhi basal",
        dibuat_pada="2026-07-31T08:00:00",
        catatan="Latihan kelas B",
    )
    tabel = DiagnosisService().rakit_tabel([
        DiagnosisPilihan("D.0008", 1, []),
        DiagnosisPilihan("D.0022", 2, ["Batasi asupan cairan dan garam sesuai program"]),
    ])
    return asesmen, tabel


def main() -> int:
    asesmen, tabel = contoh()
    service = DiagnosisService()

    print("=" * 62)
    print("TEST 1 -- Penyusun isi sel")
    print("=" * 62)
    item = tabel[0]
    sel_dx = E._sel_diagnosis(item)
    check("Sel diagnosis memuat kode & nama",
          "D.0008" in sel_dx and "Penurunan Curah Jantung" in sel_dx, sel_dx)
    check("Sel diagnosis multi-baris", "\n" in sel_dx)

    sel_luaran = E._sel_luaran(item)
    check("Sel luaran memuat kode SLKI", "L.02008" in sel_luaran, sel_luaran)

    sel_iv = E._sel_intervensi(item)
    check("Sel intervensi punya sub-judul kategori",
          "Observasi" in sel_iv and "Terapeutik" in sel_iv)
    check("Tindakan bernomor", "1. " in sel_iv and "2. " in sel_iv)
    check("Tiap tindakan pada barisnya sendiri (bukan menggumpal)",
          len(sel_iv.split("\n")) >= 10, len(sel_iv.split("\n")))
    check("Tidak memakai pemisah titik koma seperti CSV lama",
          "; Identifikasi" not in sel_iv and "; Monitor" not in sel_iv)

    print("\n" + "=" * 62)
    print("TEST 2 -- Intervensi terpilih dihormati")
    print("=" * 62)
    # Diagnosis kedua hanya mencentang 1 tindakan terapeutik.
    terapeutik = E._tindakan_terpilih(tabel[1], "terapeutik")
    check("Hanya tindakan tercentang yang muncul", len(terapeutik) == 1, terapeutik)
    # Diagnosis pertama tidak mencentang apa pun -> seluruhnya tampil.
    observasi = E._tindakan_terpilih(tabel[0], "observasi")
    semua_obs = service.intervensi("D.0008")["observasi"]
    check("Tanpa centang -> seluruh intervensi tampil sebagai acuan",
          len(observasi) == len(semua_obs), f"{len(observasi)} vs {len(semua_obs)}")

    print("\n" + "=" * 62)
    print("TEST 3 -- Markdown")
    print("=" * 62)
    md = E.ke_markdown(asesmen, tabel)
    check("Memuat nomor asesmen", asesmen.nomor in md)
    check("Memuat data S", "sesak saat berbaring" in md)
    check("Memuat data O", "JVP meningkat" in md)
    check("Memuat kedua diagnosis", "D.0008" in md and "D.0022" in md)
    check("Memuat kode luaran", "L.02008" in md)
    check("Memuat catatan asesmen", "Latihan kelas B" in md)
    check("Tabel kosong tetap menghasilkan dokumen",
          "Belum ada diagnosis" in E.ke_markdown(asesmen, []))

    print("\n" + "=" * 62)
    print("TEST 4 -- Word (.docx)")
    print("=" * 62)
    check("python-docx tersedia", E.docx_tersedia())

    data = E.ke_docx(asesmen, tabel)
    check("Menghasilkan bytes", isinstance(data, bytes) and len(data) > 5000, len(data))

    from docx import Document

    dok = Document(io.BytesIO(data))
    check("Ada tepat 1 tabel", len(dok.tables) == 1, len(dok.tables))

    t = dok.tables[0]
    check("3 baris (1 judul + 2 diagnosis)", len(t.rows) == 3, len(t.rows))
    check("4 kolom", len(t.columns) == 4, len(t.columns))
    check("Judul kolom sesuai format askep",
          [c.text for c in t.rows[0].cells] == E.JUDUL_KOLOM,
          [c.text for c in t.rows[0].cells])

    b1 = t.rows[1].cells
    check("Kolom No berisi prioritas", b1[0].text.strip() == "1", b1[0].text)
    check("Kolom diagnosis memuat kode & nama",
          "D.0008" in b1[1].text and "Penurunan Curah Jantung" in b1[1].text)
    check("Kolom luaran memuat kode SLKI", "L.02008" in b1[2].text)
    check("Kolom intervensi punya sub-judul kategori",
          "Observasi" in b1[3].text and "Terapeutik" in b1[3].text)
    check("Intervensi bernomor & multi-baris",
          b1[3].text.count("\n") >= 8, b1[3].text.count("\n"))

    bagian = dok.sections[0]
    check("Orientasi lanskap (tabel 4 kolom butuh lebar)",
          bagian.page_width > bagian.page_height)
    check("Ukuran A4 (bukan Letter bawaan python-docx)",
          abs(bagian.page_width.pt - 842) < 2 and abs(bagian.page_height.pt - 595) < 2,
          f"{bagian.page_width.pt:.0f}x{bagian.page_height.pt:.0f}")

    # Tanpa ini Word MENGABAIKAN lebar kolom dan melebarkan tabel mengikuti
    # isi terpanjang — kolom intervensi meluber keluar halaman dan
    # teksnya terpotong saat dicetak.
    from docx.oxml.ns import qn

    check("autofit dimatikan", t.autofit is False, t.autofit)
    tbl_pr = t._tbl.tblPr
    tata_letak = tbl_pr.find(qn("w:tblLayout"))
    check("Tata letak tabel dikunci 'fixed'",
          tata_letak is not None and tata_letak.get(qn("w:type")) == "fixed")
    check("Lebar tabel dinyatakan (tblW)", tbl_pr.find(qn("w:tblW")) is not None)

    tersedia = bagian.page_width.pt - bagian.left_margin.pt - bagian.right_margin.pt
    total_kolom = sum(c.width.pt for c in t.rows[0].cells)
    check("Tabel muat dalam lebar halaman",
          total_kolom <= tersedia + 1, f"{total_kolom:.0f} vs {tersedia:.0f}")
    lebar = [c.width.pt for c in t.rows[0].cells]
    check("Kolom intervensi paling lebar",
          lebar[3] > lebar[1] > lebar[2] or lebar[3] > lebar[1],
          [round(x) for x in lebar])

    teks_dok = "\n".join(p.text for p in dok.paragraphs)
    check("Data S & O ikut di dokumen",
          "sesak saat berbaring" in teks_dok and "JVP meningkat" in teks_dok)
    check("Catatan ikut di dokumen", "Latihan kelas B" in teks_dok)
    check("Tabel kosong tidak membuat gagal",
          isinstance(E.ke_docx(asesmen, []), bytes))

    print("\n" + "=" * 62)
    print("TEST 5 -- Excel (.xlsx)")
    print("=" * 62)
    check("openpyxl tersedia", E.xlsx_tersedia())

    data = E.ke_xlsx(asesmen, tabel)
    check("Menghasilkan bytes", isinstance(data, bytes) and len(data) > 3000, len(data))

    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data))
    ws = wb.active
    check("Nama lembar deskriptif", ws.title == "Asuhan Keperawatan", ws.title)

    baris_judul = next(
        (r for r in range(1, 20) if ws.cell(row=r, column=1).value == "No"), None
    )
    check("Baris judul tabel ditemukan", baris_judul is not None, baris_judul)

    judul = [ws.cell(row=baris_judul, column=c).value for c in range(1, 5)]
    check("Judul kolom sesuai format askep", judul == E.JUDUL_KOLOM, judul)

    r1 = baris_judul + 1
    check("Kolom No berisi prioritas", ws.cell(row=r1, column=1).value == 1)
    check("Kolom diagnosis terisi", "D.0008" in str(ws.cell(row=r1, column=2).value))
    check("Kolom luaran terisi", "L.02008" in str(ws.cell(row=r1, column=3).value))

    sel_iv = str(ws.cell(row=r1, column=4).value)
    check("Intervensi multi-baris di dalam sel",
          len(sel_iv.split("\n")) >= 10, len(sel_iv.split("\n")))
    check("Ada sub-judul kategori", "Observasi" in sel_iv and "Terapeutik" in sel_iv)

    check("wrap_text aktif (isi tidak terpotong)",
          ws.cell(row=r1, column=4).alignment.wrap_text is True)
    check("Tinggi baris disesuaikan isi (bukan tinggi bawaan)",
          (ws.row_dimensions[r1].height or 0) > 100, ws.row_dimensions[r1].height)

    lebar = [round(ws.column_dimensions[c].width) for c in "ABCD"]
    check("Lebar kolom diatur, kolom intervensi paling lebar",
          lebar[3] > lebar[1] > lebar[0], lebar)

    check("Diatur untuk cetak lanskap",
          str(ws.page_setup.orientation) == "landscape", ws.page_setup.orientation)
    check("Baris judul diulang tiap halaman cetak",
          bool(ws.print_title_rows), ws.print_title_rows)

    semua_teks = "\n".join(
        str(c.value) for baris in ws.iter_rows() for c in baris if c.value
    )
    check("Data S & O ikut di lembar",
          "sesak saat berbaring" in semua_teks and "JVP meningkat" in semua_teks)
    check("Catatan ikut di lembar", "Latihan kelas B" in semua_teks)

    print("\n" + "=" * 62)
    print("TEST 6 -- Nama berkas")
    print("=" * 62)
    for ext in ("docx", "xlsx", "md"):
        check(f"Nama berkas .{ext} memakai nomor asesmen",
              E.nama_berkas(asesmen, ext) == f"{asesmen.nomor}.{ext}")
    check("Tanpa nomor -> nama cadangan",
          E.nama_berkas(Asesmen(), "docx") == "asuhan.docx")

    print("\n" + "=" * 62)
    print("TEST 7 -- Ekspor CSV lama sudah tidak ada")
    print("=" * 62)
    check("Fungsi ke_csv dihapus", not hasattr(E, "ke_csv"))
    check("Digantikan ke_xlsx & ke_docx",
          hasattr(E, "ke_xlsx") and hasattr(E, "ke_docx"))

    print("\n" + "=" * 62)
    print(f"HASIL AKHIR: {PASS} PASS, {FAIL} FAIL")
    print("=" * 62)
    return 1 if FAIL else 0


if __name__ == "__main__":
    sys.exit(main())
